#!/usr/bin/env python3
"""
后处理 docx 文件：
- 中文字体：宋体
- 英文/数字：Times New Roman
- 字体颜色：黑色
- 表格：三线表
- 图片：居中对齐 + 下方图注
"""

import sys
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def is_chinese(char):
    """判断字符是否为中文"""
    return '\u4e00' <= char <= '\u9fff'


def set_run_font(run, cn_font='宋体', en_font='Times New Roman', size=Pt(12), color=RGBColor(0, 0, 0)):
    """设置 run 的字体，按字符区分中英文"""
    run.font.size = size
    run.font.color.rgb = color
    
    # 处理每个字符的字体
    text = run.text
    if not text:
        return
    
    # 如果 run 中全是英文/数字或全是中文，直接设置
    has_cn = any(is_chinese(c) for c in text)
    has_en = any(not is_chinese(c) and c.strip() for c in text)
    
    if has_cn and not has_en:
        run.font.name = cn_font
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), cn_font)
    elif has_en and not has_cn:
        run.font.name = en_font
    else:
        # 混合文本：需要拆分
        # 这里简单处理：设置西文字体为 Times New Roman，中文字体为宋体
        run.font.name = en_font
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), cn_font)


def process_paragraphs(doc):
    """处理所有段落的字体和首行缩进"""
    for para in doc.paragraphs:
        for run in para.runs:
            set_run_font(run)
        para.paragraph_format.space_after = Pt(6)
        
        # 正文段落首行缩进 2 字符（约 0.74 cm）
        style_name = para.style.name
        text = para.text.strip()
        
        # 跳过标题、表格内段落、引用、图片、图注
        if (style_name.startswith('Heading') or 
            style_name.startswith('Table') or
            style_name == 'Quote' or
            style_name == 'Block Text' or
            text.startswith('图') or
            text.startswith('上图') or
            text.startswith('下图') or
            text.startswith('>') or
            text.startswith('—') or
            text == ''):
            continue
        
        # 检查段落是否在表格中
        in_table = False
        try:
            # 通过检查祖先元素是否在表格中
            p = para._element
            # 简单判断：如果段落前后有表格相关的 xml 标记
            in_table = p.getparent().tag.endswith('tc')
        except:
            pass
        
        if in_table:
            continue
        
        # 检查段落是否包含图片
        has_image = False
        for run in para.runs:
            if 'Drawing' in run._element.xml or 'graphicData' in run._element.xml:
                has_image = True
                break
        if has_image:
            continue
        
        # 设置首行缩进 0.74 cm（约等于 2 个汉字）
        para.paragraph_format.first_line_indent = Pt(24)


def process_tables_as_three_line(doc):
    """将表格转为三线表样式"""
    for table in doc.tables:
        # 设置表格整体边框为无
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
        
        # 设置表格宽度为100%
        tblW = tblPr.find(qn('w:tblW'))
        if tblW is None:
            tblW = OxmlElement('w:tblW')
            tblPr.append(tblW)
        tblW.set(qn('w:type'), 'pct')
        tblW.set(qn('w:w'), '5000')
        
        # 设置边框
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'none')
            border.set(qn('w:sz'), '0')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'auto')
            tblBorders.append(border)
        
        # 如果 tblPr 中已有 tblBorders，先移除
        existing_borders = tblPr.find(qn('w:tblBorders'))
        if existing_borders is not None:
            tblPr.remove(existing_borders)
        tblPr.append(tblBorders)
        
        if tbl.tblPr is None:
            tbl.insert(0, tblPr)
        
        # 处理每一行
        for i, row in enumerate(table.rows):
            for cell in row.cells:
                # 设置单元格字体
                for para in cell.paragraphs:
                    for run in para.runs:
                        set_run_font(run, size=Pt(10.5))
                    para.paragraph_format.space_after = Pt(0)
                    para.paragraph_format.space_before = Pt(0)
                
                # 获取单元格属性
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                
                # 第一行（标题行）和最后一行处理底边框
                # 三线表：顶线（第一行顶）、标题行底线、底线（最后一行底）
                
        # 重新处理边框：三线表
        # 顶线：第一行顶部
        # 标题分隔线：第一行底部
        # 底线：最后一行底部
        
        if len(table.rows) > 0:
            # 清除所有单元格边框，然后设置三线表
            for row in table.rows:
                for cell in row.cells:
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    tcBorders = tcPr.find(qn('w:tcBorders'))
                    if tcBorders is not None:
                        tcPr.remove(tcBorders)
            
            # 设置第一行顶部边框（粗线）
            for cell in table.rows[0].cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                top = OxmlElement('w:top')
                top.set(qn('w:val'), 'single')
                top.set(qn('w:sz'), '12')  # 1.5pt
                top.set(qn('w:space'), '0')
                top.set(qn('w:color'), '000000')
                tcBorders.append(top)
                
                # 第一行底部边框（细线）
                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), '6')  # 0.75pt
                bottom.set(qn('w:space'), '0')
                bottom.set(qn('w:color'), '000000')
                tcBorders.append(bottom)
                
                tcPr.append(tcBorders)
            
            # 设置最后一行底部边框（粗线）
            last_row_idx = len(table.rows) - 1
            for cell in table.rows[last_row_idx].cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = tcPr.find(qn('w:tcBorders'))
                if tcBorders is None:
                    tcBorders = OxmlElement('w:tcBorders')
                    tcPr.append(tcBorders)
                
                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), '12')  # 1.5pt
                bottom.set(qn('w:space'), '0')
                bottom.set(qn('w:color'), '000000')
                tcBorders.append(bottom)


def center_images_and_add_captions(doc):
    """图片居中对齐，并处理图注"""
    for para in doc.paragraphs:
        # 检查段落是否包含图片
        has_image = False
        for run in para.runs:
            if ' graphicData ' in run._element.xml or 'Drawing' in run._element.xml:
                has_image = True
                break
        
        if has_image:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # 设置段后间距
            para.paragraph_format.space_after = Pt(6)
            
            # 处理该段落中的文字作为图注
            # 如果段落中有文本（图注文字），确保它也在图片下方
            for run in para.runs:
                if run.text.strip() and 'graphicData' not in run._element.xml:
                    set_run_font(run, size=Pt(10.5))


def process_caption_paragraphs(doc):
    """处理以 '图' 或 '上图' 开头的图注段落，使其居中并缩小字号"""
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.startswith('图') or text.startswith('上图') or text.startswith('下图'):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                set_run_font(run, size=Pt(10.5))


def set_heading_fonts(doc):
    """设置标题字体"""
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            level = para.style.name.replace('Heading', '').strip()
            try:
                level = int(level)
            except:
                level = 1
            
            size = Pt({1: 16, 2: 14, 3: 12}.get(level, 12))
            for run in para.runs:
                set_run_font(run, cn_font='黑体', en_font='Times New Roman', size=size)
                run.bold = True


def main():
    input_path = sys.argv[1] if len(sys.argv) > 1 else 'doc/report.docx'
    output_path = sys.argv[2] if len(sys.argv) > 2 else 'doc/report_formatted.docx'
    
    print(f"Processing {input_path}...")
    doc = Document(input_path)
    
    process_paragraphs(doc)
    set_heading_fonts(doc)
    process_tables_as_three_line(doc)
    center_images_and_add_captions(doc)
    process_caption_paragraphs(doc)
    
    doc.save(output_path)
    print(f"Saved to {output_path}")


if __name__ == '__main__':
    main()
